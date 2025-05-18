from pathlib import Path
import magic
import docx
import pdfplumber
from pdfplumber.utils.exceptions import PdfminerException


def extract_text(path: str) -> str:
    path = Path(path)
    mtime = path.stat().st_mtime
    print(f"Extracted {path}")
    return _extract_text_cached(str(path), mtime)


def _extract_docx_text(path: str) -> str:
    try:
        doc = docx.Document(path)
        parts = []

        # Extract body text
        parts.extend(p.text for p in doc.paragraphs if p.text.strip())

        # Extract headers and footers from all sections
        for section in doc.sections:
            header = section.header
            footer = section.footer

            parts.extend(p.text for p in header.paragraphs if p.text.strip())
            parts.extend(p.text for p in footer.paragraphs if p.text.strip())

        return "\n".join(parts)

    except Exception as e:
        print(f"DOCX read failed: {path}: {e}")
        return ""


def _extract_text_cached(path: str, mtime: float) -> str:
    suffix = Path(path).suffix.lower()

    try:
        if suffix == ".pdf":
            with pdfplumber.open(path) as pdf:
                return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

        elif suffix == ".docx":
            return _extract_docx_text(path)

        else:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()

    except PdfminerException:
        print(f"PDF read failed: {path}")
        return ""

    except Exception as e:
        print(f"Read failed: {path}: {e}")
        return ""


def is_plain_text_file(path: Path) -> bool:
    try:
        mime = magic.from_file(str(path), mime=True)
        return mime.startswith("text/")
    except Exception:
        return False
